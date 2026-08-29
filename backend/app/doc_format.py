"""把 AI（豆包/DeepSeek/Kimi…）吐出来的"带 markdown 符号但没排版"的文本，解析成结构，
再按选定的中文办公模板重建成一个排好版的 .docx。

纯本地，`python-docx` 生成，不调任何 API——就是解析 + 套样式，用户自己的文字自己排版。

两块：
  1. `parse(text)` —— 行级解析器。主吃 markdown（AI 输出的就是这个），额外认中文公文里
     常见的"一、""（一）""1."这类不带 # 的标题写法。
  2. `build_docx(blocks, template)` —— 按模板把结构写成 Word。模板只用 Windows 中文系统
     自带的字体（宋体/仿宋/黑体/楷体），换别的机器打开不会掉字。

python-docx 的坑：中文字体必须设 `w:rFonts` 的 `w:eastAsia` 属性，光设 `.font.name`
只影响西文，中文还是默认宋体。`_set_font` 里一起设了。
"""
import io
import re
from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# 解析
# ---------------------------------------------------------------------------

@dataclass
class Span:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


@dataclass
class Block:
    kind: str  # title / heading / para / bullet / ordered / quote / code / table / hr
    level: int = 0
    spans: list[Span] = field(default_factory=list)
    text: str = ""            # code 用
    rows: list[list[str]] = field(default_factory=list)  # table 用


_CN_NUM = "一二三四五六七八九十百千零两"
_RE_CN_L1 = re.compile(rf"^\s*第[{_CN_NUM}\d]+[章节篇部分]\s*")
_RE_CN_L1b = re.compile(rf"^\s*[{_CN_NUM}]+、\s*")
_RE_CN_L2 = re.compile(rf"^\s*[（(][{_CN_NUM}]+[)）]\s*")
_RE_CN_L3 = re.compile(r"^\s*\d+[.．、]\s+\S")
_RE_CN_L4 = re.compile(r"^\s*\d+[.．]\d+[.．]?\s+\S")

_RE_INLINE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*[^*]+?\*|_[^_]+?_|`[^`]+?`)")


def _parse_spans(text: str) -> list[Span]:
    spans: list[Span] = []
    for part in _RE_INLINE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            spans.append(Span(part[2:-2], bold=True))
        elif part.startswith("`") and part.endswith("`"):
            spans.append(Span(part[1:-1], code=True))
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            spans.append(Span(part[1:-1], italic=True))
        else:
            spans.append(Span(part))
    return spans or [Span(text)]


def _split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def parse(text: str, explicit_title: str | None = None) -> list[Block]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    blocks: list[Block] = []
    has_md_heading = any(re.match(r"^#{1,6}\s", ln) for ln in lines)

    i = 0
    title_used = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # 代码块
        if line.startswith("```"):
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            blocks.append(Block("code", text="\n".join(buf)))
            i = j + 1
            continue

        # 分隔线
        if re.fullmatch(r"[-*_]{3,}", line):
            blocks.append(Block("hr"))
            i += 1
            continue

        # 表格：连续的 | ... | 行
        if line.startswith("|") and "|" in line[1:]:
            rows = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                cells = _split_table_row(lines[j])
                if not re.fullmatch(r"[\s:\-|]+", lines[j].strip()):  # 跳过 |---|---| 分隔行
                    rows.append(cells)
                j += 1
            if rows:
                blocks.append(Block("table", rows=rows))
            i = j
            continue

        # 文档标题：第一个内容块、不是 markdown 标题行、较短、无句末标点 —— 当文档大标题。
        # 不受 has_md_heading 影响（豆包常见写法是首行裸标题 + 下面全是 ## 小标题）。
        if (
            not blocks
            and not title_used
            and not explicit_title
            and not line.startswith("#")
            and not re.match(r"^\s*([-*+>|]|\d+[.)、])\s", raw)
            and len(line) <= 32
            and not re.search(r"[。！？，、；：,.!?]$", line)
        ):
            blocks.append(Block("title", spans=_parse_spans(line)))
            title_used = True
            i += 1
            continue

        # markdown 标题
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            lvl = len(m.group(1))
            content = m.group(2).strip().rstrip("#").strip()
            if lvl == 1 and not title_used and not explicit_title:
                blocks.append(Block("title", spans=_parse_spans(content)))
                title_used = True
            else:
                blocks.append(Block("heading", level=min(lvl if not title_used or explicit_title else lvl - 1, 4) or 1,
                                    spans=_parse_spans(content)))
            i += 1
            continue

        # 引用
        if line.startswith(">"):
            blocks.append(Block("quote", spans=_parse_spans(line.lstrip("> ").strip())))
            i += 1
            continue

        # 无序列表
        m = re.match(r"^(\s*)[-*+]\s+(.*)", raw)
        if m:
            indent = len(m.group(1))
            blocks.append(Block("bullet", level=1 if indent < 2 else 2, spans=_parse_spans(m.group(2).strip())))
            i += 1
            continue

        # 有序列表  1. / 1) / 1、 —— 一律当列表项（不带 # 时的"小标题"靠 一、/（一）识别，
        # 阿拉伯数字开头的行绝大多数是罗列点，硬判成标题反而容易错）
        m = re.match(r"^(\s*)\d+[.)、]\s+(.*)", raw)
        if m:
            blocks.append(Block("ordered", level=1, spans=_parse_spans(m.group(2).strip())))
            i += 1
            continue

        # 中文公文式标题（不带 #）
        if not has_md_heading:
            if _RE_CN_L1.match(line):
                blocks.append(Block("heading", level=1, spans=_parse_spans(line)))
                i += 1
                continue
            if _RE_CN_L1b.match(line) and len(line) <= 40:
                blocks.append(Block("heading", level=1, spans=_parse_spans(line)))
                i += 1
                continue
            if _RE_CN_L2.match(line) and len(line) <= 40:
                blocks.append(Block("heading", level=2, spans=_parse_spans(line)))
                i += 1
                continue
            if _RE_CN_L4.match(line):
                blocks.append(Block("heading", level=4, spans=_parse_spans(line)))
                i += 1
                continue

        blocks.append(Block("para", spans=_parse_spans(line)))
        i += 1

    if explicit_title:
        blocks.insert(0, Block("title", spans=_parse_spans(explicit_title.strip())))

    return blocks


# ---------------------------------------------------------------------------
# 生成 docx
# ---------------------------------------------------------------------------

# 只用 Windows 中文系统必有的字体，换机器打开不掉字
_TEMPLATES: dict[str, dict] = {
    "general": {
        "label": "通用文档",
        "margins": (2.54, 2.54, 3.18, 3.18),  # 上下左右 cm
        "title": {"font": "黑体", "size": 22, "align": "center", "bold": True, "after": 18},
        "headings": [
            {"font": "黑体", "size": 16, "bold": False, "before": 12, "after": 6},
            {"font": "黑体", "size": 14, "bold": False, "before": 10, "after": 4},
            {"font": "黑体", "size": 12, "bold": True, "before": 8, "after": 4},
            {"font": "宋体", "size": 12, "bold": True, "before": 6, "after": 2},
        ],
        "body": {"font": "宋体", "size": 12, "line": 1.5, "indent": True},
    },
    "report": {
        "label": "工作报告",
        "margins": (2.54, 2.54, 2.8, 2.8),
        "title": {"font": "方正小标宋简体", "size": 22, "align": "center", "bold": True, "after": 24,
                  "fallback": "宋体"},
        "headings": [
            {"font": "黑体", "size": 16, "bold": False, "before": 12, "after": 6},
            {"font": "楷体", "size": 16, "bold": True, "before": 10, "after": 4},
            {"font": "仿宋", "size": 16, "bold": True, "before": 6, "after": 2},
            {"font": "仿宋", "size": 16, "bold": True, "before": 4, "after": 2},
        ],
        "body": {"font": "仿宋", "size": 16, "line": 1.5, "indent": True},
    },
    "official": {
        # 近似 GB/T 9704-2012 党政机关公文格式：正文仿宋三号，一级"一、"黑体，
        # 二级"（一）"楷体，三级"1."仿宋加粗。字体名用系统自带的。
        "label": "公文格式",
        "margins": (3.7, 3.5, 2.8, 2.6),
        "title": {"font": "方正小标宋简体", "size": 22, "align": "center", "bold": True, "after": 24,
                  "fallback": "宋体"},
        "headings": [
            {"font": "黑体", "size": 16, "bold": False, "before": 0, "after": 0},
            {"font": "楷体", "size": 16, "bold": False, "before": 0, "after": 0},
            {"font": "仿宋", "size": 16, "bold": True, "before": 0, "after": 0},
            {"font": "仿宋", "size": 16, "bold": True, "before": 0, "after": 0},
        ],
        "body": {"font": "仿宋", "size": 16, "line": 1.5625, "indent": True},  # 28磅 / 16pt ≈ 1.5625 行距近似
    },
}


def templates() -> list[dict]:
    return [{"key": k, "label": v["label"]} for k, v in _TEMPLATES.items()]


def _set_font(run, cjk_name: str, size_pt: float, bold: bool = False, italic: bool = False):
    """cjk_name 是中文字体名。西文/数字统一走 Times New Roman（等宽代码走 Consolas），
    符合 GB/T 9704「数字和字母用 Times New Roman」，也避免中英混排字体不一致。"""
    latin = "Consolas" if cjk_name == "Consolas" else "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), cjk_name)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def _add_spans(paragraph, spans: list[Span], font: str, size: float, base_bold: bool = False):
    for sp in spans:
        run = paragraph.add_run(sp.text)
        _set_font(
            run,
            "Consolas" if sp.code else font,
            size,
            bold=base_bold or sp.bold,
            italic=sp.italic,
        )


def build_docx(blocks: list[Block], template_key: str) -> bytes:
    tpl = _TEMPLATES.get(template_key, _TEMPLATES["general"])
    doc = Document()

    top, bottom, left, right = tpl["margins"]
    from docx.shared import Cm

    for section in doc.sections:
        section.top_margin, section.bottom_margin = Cm(top), Cm(bottom)
        section.left_margin, section.right_margin = Cm(left), Cm(right)

    normal = doc.styles["Normal"]
    normal.font.name = tpl["body"]["font"]
    normal.font.size = Pt(tpl["body"]["size"])
    normal.paragraph_format.line_spacing = tpl["body"]["line"]

    body_f, body_s = tpl["body"]["font"], tpl["body"]["size"]

    for b in blocks:
        if b.kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if tpl["title"]["align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(tpl["title"]["after"])
            _add_spans(p, b.spans, tpl["title"].get("fallback", tpl["title"]["font"]), tpl["title"]["size"],
                       base_bold=tpl["title"]["bold"])

        elif b.kind == "heading":
            lvl = min(max(b.level, 1), len(tpl["headings"]))
            hs = tpl["headings"][lvl - 1]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(hs["before"])
            p.paragraph_format.space_after = Pt(hs["after"])
            p.paragraph_format.line_spacing = tpl["body"]["line"]
            if template_key == "official":
                p.paragraph_format.first_line_indent = Pt(hs["size"])  # 公文标题也缩进
            else:
                p.paragraph_format.first_line_indent = Pt(0)
            # 挂大纲级别，Word 的导航窗格/目录才认得出这是标题
            pr = p._p.get_or_add_pPr()
            ol = pr.makeelement(qn("w:outlineLvl"), {})
            ol.set(qn("w:val"), str(lvl - 1))
            pr.append(ol)
            _add_spans(p, b.spans, hs["font"], hs["size"], base_bold=hs["bold"])

        elif b.kind == "para":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = tpl["body"]["line"]
            if tpl["body"]["indent"]:
                p.paragraph_format.first_line_indent = Pt(body_s * 2)
            _add_spans(p, b.spans, body_f, body_s)

        elif b.kind in ("bullet", "ordered"):
            style = "List Bullet" if b.kind == "bullet" else "List Number"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
            p.paragraph_format.line_spacing = tpl["body"]["line"]
            _add_spans(p, b.spans, body_f, body_s)

        elif b.kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(body_s * 2)
            p.paragraph_format.line_spacing = tpl["body"]["line"]
            _add_spans(p, b.spans, "楷体", body_s)

        elif b.kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(body_s)
            for k, ln in enumerate(b.text.split("\n")):
                if k:
                    run = p.add_run()
                    run.add_break()
                run = p.add_run(ln)
                _set_font(run, "Consolas", body_s - 1)

        elif b.kind == "hr":
            p = doc.add_paragraph()
            pr = p._p.get_or_add_pPr()
            bd = pr.makeelement(qn("w:pBdr"), {})
            bottom_b = pr.makeelement(qn("w:bottom"), {})
            for key, val in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "999999")):
                bottom_b.set(qn(key), val)
            bd.append(bottom_b)
            pr.append(bd)

        elif b.kind == "table" and b.rows:
            cols = max(len(r) for r in b.rows)
            t = doc.add_table(rows=len(b.rows), cols=cols)
            t.style = "Table Grid"
            for ri, row in enumerate(b.rows):
                for ci in range(cols):
                    cell = t.cell(ri, ci)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    _add_spans(para, _parse_spans(row[ci] if ci < len(row) else ""), body_f, body_s - 1,
                               base_bold=(ri == 0))
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def format_to_docx(text: str, template_key: str, title: str | None = None) -> bytes:
    blocks = parse(text, explicit_title=title or None)
    if not blocks:
        raise ValueError("没有可排版的内容")
    return build_docx(blocks, template_key)
